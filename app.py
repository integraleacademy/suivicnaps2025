
from flask import Flask, render_template, request, redirect, send_file
import sqlite3
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
DB_NAME = "cnaps.db"

def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS dossiers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nom TEXT NOT NULL,
                prenom TEXT NOT NULL,
                formation TEXT NOT NULL,
                session TEXT NOT NULL,
                lien TEXT,
                statut TEXT DEFAULT 'INCOMPLET', commentaire TEXT
            )
        """)

@app.route("/")
def index():
    with sqlite3.connect(DB_NAME) as conn:
        dossiers = conn.execute("SELECT * FROM dossiers").fetchall()
    return render_template("index.html", dossiers=dossiers)

@app.route("/add", methods=["POST"])
def add():
    nom = request.form["nom"]
    prenom = request.form["prenom"]
    formation = request.form["formation"]
    session = request.form["session"]
    lien = request.form["lien"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("INSERT INTO dossiers (nom, prenom, formation, session, lien, statut) VALUES (?, ?, ?, ?, ?, ?)",
                     (nom, prenom, formation, session, lien, "INCOMPLET"))
    return redirect("/")

@app.route("/edit/<int:id>", methods=["POST"])
def edit(id):
    lien = request.form["lien"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET lien = ? WHERE id = ?", (lien, id))
    return redirect("/")

@app.route("/statut/<int:id>/<string:new_status>")
def update_statut(id, new_status):
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET statut = ? WHERE id = ?", (new_status, id))
    return redirect("/")

@app.route("/delete/<int:id>")
def delete(id):
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("DELETE FROM dossiers WHERE id = ?", (id,))
    return redirect("/")

@app.route("/attestation/<int:id>")
def attestation(id):
    with sqlite3.connect(DB_NAME) as conn:
        dossier = conn.execute("SELECT * FROM dossiers WHERE id = ?", (id,)).fetchone()

    nom, prenom, formation, session = dossier[1], dossier[2], dossier[3], dossier[4]
    date_today = datetime.now().strftime("%d/%m/%Y")

    doc = Document()
    doc.add_heading("ANNEXE : Justificatif de préinscription à une formation", 0)
    doc.add_paragraph("Cadre réservé à l’organisme de formation")
    doc.add_paragraph(f"Je soussigné(e), Monsieur Clément VAILLANT")
    doc.add_paragraph("Responsable de l'organisme de formation : INTEGRALE SECURITE FORMATIONS")
    doc.add_paragraph("Numéro d'enregistrement DIRECCTE : 93830600283")
    doc.add_paragraph("Autorisé à exercer par le CNAPS sous le numéro : FOR-083-2027-02-08-20220755135")
    doc.add_paragraph("Téléphone : 04 22 47 07 68")
    doc.add_paragraph("Adresse électronique : integralesecuriteformations@gmail.com")
    doc.add_paragraph(f"Certifie que Monsieur / Madame : {nom} {prenom}")
    doc.add_paragraph("est préinscrit(e) à la formation qualifiante ci-dessous :")

    if formation == "A3P":
        doc.add_paragraph("Libellé exact de la formation : AGENT DE PROTECTION PHYSIQUE DES PERSONNES (A3P)")
        doc.add_paragraph("Numéro d'enregistrement RNCP : 35098")
        doc.add_paragraph("Nature de la formation : Titre à Finalité Professionnelle (TFP) Agent de Protection Physique des Personnes - Agrément de la CPNEFP n°8320111201 en date du 02/02/2021")
    else:
        doc.add_paragraph("Libellé exact de la formation : AGENT DE PREVENTION ET DE SECURITE (APS)")
        doc.add_paragraph("Numéro d'enregistrement RNCP : 34054")
        doc.add_paragraph("Nature de la formation : Titre à Finalité Professionnelle (TFP) Agent de Prévention et de Sécurité - Agrément de la CPNEFP n°8320032701 en date du 30/11/2020")

    doc.add_paragraph(f"Dates de la formation : {session} qui se déroulera à Puget sur Argens (83480).")
    doc.add_paragraph("Lieu(x) de réalisation de la formation : Intégrale Sécurité Formations - 54 chemin du Carreou - 83480 PUGET SUR ARGENS")
    doc.add_paragraph("")
    doc.add_paragraph("Monsieur Clément VAILLANT")
    doc.add_paragraph("Directeur Général – Intégrale Sécurité Formations")

    doc.add_picture('static/signature_bloc.png', width=Inches(3.8))

    filename = f"attestation_{nom}_{prenom}.docx"
    filepath = os.path.join("temp", filename)
    os.makedirs("temp", exist_ok=True)
    doc.save(filepath)

    return send_file(filepath, as_attachment=True)

if __name__ == "__main__":
    init_db()
    port = int(os.environ.get("PORT", 10000))
    app.run(debug=True, host="0.0.0.0", port=port)

@app.route("/commentaire/<int:id>", methods=["POST"])
def commentaire(id):
    texte = request.form["commentaire"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET commentaire = ? WHERE id = ?", (texte, id))
    return redirect("/")
